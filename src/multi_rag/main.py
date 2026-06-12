import pymupdf
from PIL import Image
import io
import os
import re
from sentence_transformers import SentenceTransformer
from chromadb import PersistentClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
import json
import base64
from langchain_core.messages import HumanMessage
from langchain.chat_models import init_chat_model
import zipfile
import docx
import pandas as pd

model=init_chat_model(model="gemini-2.5-flash",model_provider="google-genai")
image_embedder=SentenceTransformer("clip-ViT-B-32")
text_embedder=SentenceTransformer("BAAI/bge-base-en-v1.5") #"BAAI/bge-m3"

client=PersistentClient(path="./chroma_db")
text_store = client.get_or_create_collection("rag-text",configuration={"hnsw":{"space":"cosine"}})
image_store = client.get_or_create_collection("rag-image",configuration={"hnsw":{"space":"ip"}})
table_store=client.get_or_create_collection("rag-table",configuration={"hnsw":{"space":"cosine"}})
text_splitter=RecursiveCharacterTextSplitter(chunk_size=1000,chunk_overlap=100)
os.makedirs("./temp", exist_ok=True)
os.makedirs("./temp/images", exist_ok=True)

def docnamer(docname: str):
    ref_docname = docname
    ref_docname=re.sub(r"[<>:/\\|?*]",'.',ref_docname)
    ref_docname=ref_docname.replace(' ','-')
    ref_docname=ref_docname.lstrip('.')
    return ref_docname

def embed(docname:str):
    ext=docname[docname.rfind('.')+1:]
    if ext=="docx" or ext=="doc":
        embed_from_doc(docname)
    elif ext=="pdf":
        embed_from_pdf(docname)
    elif ext=="txt":
        embed_from_txt(docname)
    elif ext=="png" or ext=="jpeg" or ext=="jpg":
        embed_from_image(docname)
    elif ext=="xlsx":
        embed_from_xlsx(docname)

def embed_from_pdf(docname:str):
    ref_docname=docnamer(docname)
    doc=pymupdf.open(docname)
    tmd=[]
    imd=[]
    tamd=[]
    images=[]
    tids=[]
    iids=[]
    taids=[]
    tables=[]
    splits=[]
    pages=0
    for page in doc:
        text=page.get_text()
        img_list=page.get_images()
        tab_list=page.find_tables()
        os.makedirs(f"./temp/images/{ref_docname}", exist_ok=True)
        if tab_list:
            for table in tab_list.tables:
                table=table.to_markdown()
                tables.append(table)
            tamd.extend([{"docname": f"{ref_docname}","page": f"{pages}","table": f"{i}","type": "table"} for i in range(len(tab_list.tables))])
            taids.extend([f"{ref_docname}_p{pages}_t{i}" for i in range(len(tab_list.tables))])
        if img_list:
            for img in img_list:
                img=doc.extract_image(img[0])
                img=Image.open(io.BytesIO(img['image']))
                images.append(img)
            imd.extend([{"docname": f"{ref_docname}","page": f"{pages}","imno": f"{i}","type": "image/png"} for i in range(len(img_list))])
            iids.extend([f"{ref_docname}_p{pages}_i{i}" for i in range(len(img_list))])
        split=text_splitter.split_text(text)
        tids.extend([f"{ref_docname}_p{pages}_c{i}" for i in range(len(split))])
        tmd.extend([{"docname": f"{ref_docname}","page": f"{pages}","chunk": f"{i}","type": "text",} for i in range(len(split))])
        splits.extend(split)
        pages+=1

    for iid,image in zip(iids,images):
        image.save(f"./temp/images/{ref_docname}/{iid}.png")
    if os.path.exists("./temp/text.json"):
        with open("./temp/text.json", "r", encoding="utf-8") as f:
            text_lookup = json.load(f)
    else:
        text_lookup = {}
    text_lookup.setdefault(ref_docname,{})
    text_lookup[ref_docname].update(dict(zip(tids,splits)))
    with open(f"./temp/text.json","w",encoding="utf-8") as f:
        json.dump(text_lookup,f,indent=4)
    
    if os.path.exists("./temp/tables.json"):
        with open("./temp/tables.json", "r", encoding="utf-8") as f:
            table_lookup = json.load(f)
    else:
        table_lookup = {}
    table_lookup.setdefault(ref_docname,{})
    table_lookup[ref_docname].update(dict(zip(taids,tables)))
    with open(f"./temp/tables.json","w",encoding="utf-8") as f:
        json.dump(table_lookup,f,indent=4)

    if splits:
        t_emb=text_embedder.encode(splits)
        text_store.upsert(ids=tids,embeddings=t_emb,metadatas=tmd)
    if images:
        i_emb=image_embedder.encode(images)
        image_store.upsert(ids=iids,embeddings=i_emb,metadatas=imd)
    if tables:
        ta_emb=text_embedder.encode(tables)
        table_store.upsert(ids=taids,embeddings=ta_emb,metadatas=tamd)

def embed_from_doc(docname: str):
    ref_docname = docnamer(docname)
    os.makedirs(f"./temp/images/{ref_docname}", exist_ok=True)
    tmd=[]
    imd=[]
    tamd=[]
    images = []
    tids = []
    iids = []
    taids = []
    tables = []
    splits = []

    # 1. EXTRACT AND SAVE IMAGES FROM ZIP ARCHIVE
    i = 0
    with zipfile.ZipFile(docname, "r") as archive:
        for file in archive.namelist():
            if file.startswith("word/media/"):
                ext = file.split(".")[-1]
                # Open image and force load data into RAM immediately
                img = Image.open(io.BytesIO(archive.read(file)))
                img.load()
                images.append(img)

                img_id = f"{ref_docname}_i{i}"
                iids.append(img_id)
                imd.append({"docname":f"{ref_docname}","imno":f"{i}","type":f"image/{ext}"})

                # Save the image right away to your temp folder
                img.save(f"./temp/images/{ref_docname}/{img_id}.{ext}")
                i += 1

    # 2. EXTRACT TEXT AND CONVERT TABLES TO MARKDOWN STRINGS
    doc = docx.Document(docname)

    paragraph_index = 0
    for p in doc.paragraphs:
        text_content = p.text.strip()
        if text_content:
            split = text_splitter.split_text(text_content)
            splits.extend(split)
            tids.extend([f"{ref_docname}_p{paragraph_index}_c{chunk_index}" for chunk_index in range(len(split))])
            tmd.extend([{"docname":f"{ref_docname}","pno":f"{paragraph_index}","chunk":f"{chunk_index}","type":"text"} for chunk_index in range(len(split))])
            paragraph_index += 1

    # Process Document Tables (and convert to Markdown strings)
    for table_idx, table in enumerate(doc.tables):
        markdown_rows = []
        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]

            # Generate Markdown format string grid
            markdown_rows.append("| " + " | ".join(row_cells) + " |")

            # Insert Markdown header separator line under the first header row
            if row_idx == 0:
                separator = "|" + "|".join(["---"] * len(row_cells)) + "|"
                markdown_rows.append(separator)

        table_markdown = "\n".join(markdown_rows)
        tables.append(table_markdown)
        taids.append(f"{ref_docname}_t{table_idx}")
        tamd.append({"docname":f"{ref_docname}","table":f"{table_idx}","type":"table"})

    if os.path.exists("./temp/text.json"):
        with open("./temp/text.json", "r", encoding="utf-8") as f:
            text_lookup = json.load(f)
    else:
        text_lookup = {}
    text_lookup.setdefault(ref_docname, {})
    text_lookup[ref_docname].update(dict(zip(tids, splits)))
    with open("./temp/text.json", "w", encoding="utf-8") as f:
        json.dump(text_lookup, f, indent=4)

    if os.path.exists("./temp/tables.json"):
        with open("./temp/tables.json", "r", encoding="utf-8") as f:
            table_lookup = json.load(f)
    else:
        table_lookup = {}
    table_lookup.setdefault(ref_docname, {})
    table_lookup[ref_docname].update(dict(zip(taids, tables)))
    with open("./temp/tables.json", "w", encoding="utf-8") as f:
        json.dump(table_lookup, f, indent=4)

    if splits:
        t_emb = text_embedder.encode(splits)
        text_store.upsert(ids=tids, embeddings=t_emb,metadatas=tmd)
    if images:
        i_emb = image_embedder.encode(images)
        image_store.upsert(ids=iids, embeddings=i_emb,metadatas=imd)
    if tables:
        ta_emb = text_embedder.encode(tables)
        table_store.upsert(ids=taids, embeddings=ta_emb,metadatas=tamd)

def embed_from_txt(docname:str):
    ref_docname=docnamer(docname)
    with open(docname,"r",encoding="utf-8") as f:
        text=f.read()
    splits=text_splitter.split_text(text)
    ids=[f"{ref_docname}_c{i}" for i in range(len(splits))]
    mds=[{"docname":ref_docname,"chunk":i,"type":"text"} for i in range(len(splits))]
    emb=text_embedder.encode(splits)
    text_store.upsert(ids=ids,metadatas=mds,embeddings=emb)
    if os.path.exists("./temp/text.json"):
        with open("./temp/text.json", "r", encoding="utf-8") as f:
            text_lookup = json.load(f)
    else:
        text_lookup = {}
    text_lookup.setdefault(ref_docname, {})
    text_lookup[ref_docname].update(dict(zip(ids, splits)))
    with open("./temp/text.json", "w", encoding="utf-8") as f:
        json.dump(text_lookup, f, indent=4)

def embed_from_image(docname:str):
    ref_docname=docnamer(docname)
    img=Image.open(docname)
    emb=image_embedder.encode([img])
    id=[f"{ref_docname}_0"]
    md=[{"docname":f"{ref_docname}","imno":0,"type":f"image/{docname[docname.rfind('.')+1:]}"}]
    image_store.upsert(embeddings=emb,ids=id,metadatas=md)
    os.makedirs(f"./temp/images/{ref_docname}",exist_ok=True)
    img.save(f"./temp/images/{ref_docname}/{ref_docname}")

def embed_from_xlsx(docname:str):
    ref_docname=docnamer(docname)
    df=pd.read_excel(docname)
    tmd=[df.to_markdown()]
    md=[{"docname":ref_docname,"type":"table","table":"table"}]
    emb=text_embedder.encode(tmd)
    id=[f"{ref_docname}_0"]
    table_store.upsert(embeddings=emb,ids=id,metadatas=md)
    if os.path.exists("./temp/text.json"):
        with open("./temp/text.json", "r", encoding="utf-8") as f:
            text_lookup = json.load(f)
    else:
        text_lookup = {}
    text_lookup.setdefault(ref_docname, {})
    text_lookup[ref_docname].update(dict(zip(id, tmd)))
    with open("./temp/text.json", "w", encoding="utf-8") as f:
        json.dump(text_lookup, f, indent=4)

def retrieve(query:str):
    with open("./temp/text.json", "r", encoding="utf-8") as f:
        text_lookup = json.load(f)
    with open("./temp/tables.json", "r", encoding="utf-8") as f:
        table_lookup = json.load(f)
    qt_emb = text_embedder.encode(query)
    qi_emb=image_embedder.encode(query)

    text_results = text_store.query(
        query_embeddings=[qt_emb],
        n_results=5
    )
    image_results = image_store.query(
        query_embeddings=[qi_emb],
        n_results=2
    )
    table_results=table_store.query(
        query_embeddings=[qt_emb],
        n_results=3
    )

    text_ids = text_results["ids"][0]
    chunks=[text_lookup.get(text_results["metadatas"][0][i]["docname"],{}).get(text_ids[i],"NOT FOUND") for i in range(len(text_ids))]
    image_ids = image_results["ids"][0]
    images=[]
    for i in range(len(image_ids)):
        with open(f"./temp/images/{image_results['metadatas'][0][i]['docname']}/{image_ids[i]}.{image_results['metadatas'][0][i]['type'][6:]}","rb") as f:
            images.append(base64.b64encode(f.read()).decode("utf-8"))
    table_ids=table_results["ids"][0]
    tables=[table_lookup.get(table_results["metadatas"][0][i]["docname"],{}).get(table_ids[i],"NOT FOUND") for i in range(len(table_ids))]
    return {
        "images":images,
        "text":chunks,
        "tables":tables
    }

def query(query:str):
    result=retrieve(query)
    parts=[
    {
        "type":"text",
        "text":f"""
            Answer the Question based on the Context given.
            If the given Context doesn't contain any information regarding the Question, Answer 'Out of Context'

            Question:{query}"""
    }
    ]
    for text in result["text"]:
        parts.append({
            "type":"text",
            "text":text
        })
    for table in result["tables"]:
        parts.append({
            "type":"text",
            "text":f"TABLE:\n{table}"
        })
    for image in result["images"]:
        parts.append({
            "type":"image_url",
            "image_url":f"data:image/png;base64,{image}"
        })

    message=HumanMessage(content=parts)
    response=model.invoke([message])
    return [result,response.content]